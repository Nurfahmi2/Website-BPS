from flask import Flask, render_template, request, send_file
from docx import Document
import os

app = Flask(__name__)

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/generate', methods=['POST'])
def generate():
    # Ambil data dari form
    nomor_surat = request.form['nomor_surat']
    tanggal_surat = request.form['tanggal_surat']
    nama_petugas = request.form['nama_petugas']
    periode = request.form['periode']

    tanggal_kegiatan_list = request.form.getlist('tanggal_kegiatan[]')
    uraian_kegiatan_list = request.form.getlist('uraian_kegiatan[]')
    permasalahan_list = request.form.getlist('permasalahan[]')
    pemecahan_masalah_list = request.form.getlist('pemecahan_masalah[]')
    keterangan_list = request.form.getlist('keterangan[]')

    # Buat daftar nomor otomatis
    nomor_list = [str(i + 1) for i in range(len(tanggal_kegiatan_list))]

    # Path file template (lokal)
    template_path = os.path.join('templates', 'template.docx')
    if not os.path.exists(template_path):
        return "Template file not found. Please upload the template."

    # Buka file template
    doc = Document(template_path)

    # Temukan tabel kegiatan yang relevan
    table_kegiatan = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.rows[0].cells) >= 6:  # Tabel yang memiliki header dengan 6 kolom
            if "No." in table.rows[0].cells[0].text and "Tanggal" in table.rows[0].cells[1].text:
                table_kegiatan = table
                break

    if table_kegiatan is not None:
        # Tambahkan data kegiatan ke tabel kegiatan
        for i in range(len(tanggal_kegiatan_list)):
            if i + 2 >= len(table_kegiatan.rows):  # Jika baris kurang, tambahkan baris baru
                new_row = table_kegiatan.add_row()
            else:  # Jika baris sudah ada, gunakan baris tersebut
                new_row = table_kegiatan.rows[i + 2]

            if len(new_row.cells) >= 6:  # Validasi jumlah kolom minimal 6
                new_row.cells[0].text = nomor_list[i]  # Nomor urut
                new_row.cells[1].text = tanggal_kegiatan_list[i]
                new_row.cells[2].text = uraian_kegiatan_list[i]
                new_row.cells[3].text = permasalahan_list[i]
                new_row.cells[4].text = pemecahan_masalah_list[i]
                new_row.cells[5].text = keterangan_list[i]

        # Hapus baris kosong tambahan jika ada
        while len(table_kegiatan.rows) > len(tanggal_kegiatan_list) + 2:
            table_kegiatan._tbl.remove(table_kegiatan.rows[-1]._tr)

    # Isi informasi di bagian atas dokumen (tidak memengaruhi tabel lampiran)
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:  # Validasi jumlah kolom minimal 2
                if 'Nomor Surat Tugas' in row.cells[0].text:
                    row.cells[1].text = nomor_surat
                elif 'Tanggal Surat Tugas' in row.cells[0].text:
                    row.cells[1].text = tanggal_surat
                elif 'Nama Petugas' in row.cells[0].text:
                    row.cells[1].text = nama_petugas
                elif 'Periode Penugasan' in row.cells[0].text:
                    row.cells[1].text = periode

    # Simpan dokumen
    filename = 'Laporan_Pendataan_Lapangan.docx'
    doc.save(filename)

    # Kirim file ke pengguna
    return send_file(filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
